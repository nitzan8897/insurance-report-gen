# src/document/report_generator.py
import tkinter as tk
from tkinter import ttk, filedialog
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
from .document_utils import DocumentUtils

class ReportGenerator:
    def __init__(self):
        self.doc_utils = DocumentUtils()

    def get_safe_value(self, form_data, key, default=''):
        """
        Safely get a value from form_data, converting to string.
        Handles different widget types appropriately.
        """
        try:
            value = form_data.get(key, default)
            
            if value is None:
                return default
                
            if isinstance(value, tk.Text):
                text_content = value.get('1.0', 'end-1c')
                return text_content.strip() or default
                
            if isinstance(value, (ttk.Entry, ttk.Combobox)):
                entry_content = value.get()
                return entry_content or default
                
            if hasattr(value, 'get_date'):
                try:
                    return value.get_date().strftime('%d/%m/%Y')
                except:
                    return default
                    
            return str(value)

        except Exception as e:
            print(f"Error getting value for {key}: {str(e)}")
            return default

    def generate_header(self, doc, form_data):
        """Generate the header section of the report."""
        try:
            # Current date and reference number
            current_date = datetime.now().strftime("%d.%m.%Y")
            ref_number = datetime.now().strftime("%d%m%y")
            
            # Add date and reference
            date_para = self.doc_utils.make_hebrew_paragraph(
                doc, 
                f"תאריך: {current_date}",
                bold=True,
                alignment=WD_PARAGRAPH_ALIGNMENT.RIGHT
            )
            
            ref_para = self.doc_utils.make_hebrew_paragraph(
                doc, 
                f"מספר תיק: {ref_number}",
                bold=True,
                alignment=WD_PARAGRAPH_ALIGNMENT.RIGHT
            )

            # Add spacing
            spacing = doc.add_paragraph()
            spacing.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

            # Add recipient section
            self.doc_utils.make_hebrew_paragraph(doc, "לכבוד", bold=True)
            self.doc_utils.make_hebrew_paragraph(doc, 'הפניקס חברה לביטוח בע"מ', bold=True)

            # Add spacing
            spacing = doc.add_paragraph()
            spacing.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

            # Add report title and details
            headers = [
                "הנדון: דו\"ח חקירה",
                "======================",
                f"שם המבוטח: {self.get_safe_value(form_data, 'full_name')}",
                f"סוג האירוע: {self.get_safe_value(form_data, 'event_type')}",
                f"מספר רישוי: {self.get_safe_value(form_data, 'vehicle_license_number')}",
                f"תאריך אירוע: {self.get_safe_value(form_data, 'event_date')}",
                f"מספר תביעה: {self.get_safe_value(form_data, 'claim_number')}",
                "========================="
            ]

            for text in headers:
                self.doc_utils.make_hebrew_paragraph(
                    doc,
                    text,
                    bold=True,
                    alignment=WD_PARAGRAPH_ALIGNMENT.CENTER
                )

            # Add spacing after headers
            spacing = doc.add_paragraph()
            spacing.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        except Exception as e:
            print(f"Error in generate_header: {str(e)}")
            raise

    def generate_general_section(self, doc, form_data):
        """Generate the general section of the report."""
        try:
            # Add section title
            self.doc_utils.create_section_header(doc, "1. כללי")

            # Create general content
            vehicle_info = (
                f"נתבקשנו על ידי חברתכם לבצע חקירה בעקבות הודעת המבוטח על "
                f"{self.get_safe_value(form_data, 'event_type')} שארע/ה לו ברכבו מסוג "
                f"{self.get_safe_value(form_data, 'vehicle_company')} "
                f"{self.get_safe_value(form_data, 'vehicle_model')} "
                f"בצבע {self.get_safe_value(form_data, 'vehicle_color')}, "
                f"שנת ייצור {self.get_safe_value(form_data, 'vehicle_manufacture_year')}."
            )

            self.doc_utils.make_hebrew_paragraph(doc, vehicle_info)

            # Add investigation actions header
            self.doc_utils.make_hebrew_paragraph(
                doc, 
                "במסגרת החקירה ביצענו את הפעולות הבאות:",
                bold=True
            )

            # Add bullet points for actions taken
            actions = [
                f"פגשנו וחקרנו את המבוטח {self.get_safe_value(form_data, 'full_name')}.",
                "ערכנו בדיקה במאגרי המידע הרלוונטיים.",
                "בדקנו את מסמכי הביטוח והרישוי.",
                "צילמנו תמונות של הרכב והנזקים."
            ]

            for action in actions:
                self.doc_utils.add_bullet_point(doc, action)

            # Add spacing after general section
            spacing = doc.add_paragraph()
            spacing.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        except Exception as e:
            print(f"Error in generate_general_section: {str(e)}")
            raise

    def generate_vehicle_section(self, doc, form_data):
        """Generate the vehicle details section of the report."""
        try:
            self.doc_utils.create_section_header(doc, "2. פרטי הרכב")

            vehicle_details = [
                f"יצרן ודגם: {self.get_safe_value(form_data, 'vehicle_company')} {self.get_safe_value(form_data, 'vehicle_model')}",
                f"צבע: {self.get_safe_value(form_data, 'vehicle_color')}",
                f"שנת ייצור: {self.get_safe_value(form_data, 'vehicle_manufacture_year')}",
                f"מספר רישוי: {self.get_safe_value(form_data, 'vehicle_license_number')}",
                f"נפח מנוע: {self.get_safe_value(form_data, 'vehicle_engine_capacity')} סמ\"ק",
                f"סוג תיבת הילוכים: {self.get_safe_value(form_data, 'vehicle_gearbox')}"
            ]

            for detail in vehicle_details:
                self.doc_utils.make_hebrew_paragraph(doc, detail)

            # Add spacing after vehicle section
            spacing = doc.add_paragraph()
            spacing.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        except Exception as e:
            print(f"Error in generate_vehicle_section: {str(e)}")
            raise

    def generate_circumstances_section(self, doc, form_data):
        """Generate the circumstances section of the report."""
        try:
            circumstances = self.get_safe_value(form_data, 'circumstances')
            if circumstances:
                self.doc_utils.create_section_header(doc, "3. נסיבות האירוע")
                self.doc_utils.make_hebrew_paragraph(doc, circumstances)
                spacing = doc.add_paragraph()
                spacing.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        except Exception as e:
            print(f"Error in generate_circumstances_section: {str(e)}")
            raise

    def generate_summary_section(self, doc, form_data):
        """Generate the summary section of the report."""
        try:
            summary = self.get_safe_value(form_data, 'summary')
            if summary:
                self.doc_utils.create_section_header(doc, "4. סיכום")
                self.doc_utils.make_hebrew_paragraph(doc, summary)
                spacing = doc.add_paragraph()
                spacing.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        except Exception as e:
            print(f"Error in generate_summary_section: {str(e)}")
            raise

    def generate_signature(self, doc):
        """Generate the signature section of the report."""
        try:
            spacing = doc.add_paragraph()
            spacing.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            self.doc_utils.make_hebrew_paragraph(doc, "בכבוד רב,")
            self.doc_utils.make_hebrew_paragraph(doc, "אניגמה חקירות")

        except Exception as e:
            print(f"Error in generate_signature: {str(e)}")
            raise

    def generate(self, form_data):
        """Main method to generate the complete report."""
        try:
            doc = Document()
            
            # Configure document
            self.doc_utils.set_document_rtl(doc)
            self.doc_utils.add_custom_header_footer(doc)
            
            # Set margins
            section = doc.sections[0]
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

            # Generate all sections
            self.generate_header(doc, form_data)
            self.generate_general_section(doc, form_data)
            self.generate_vehicle_section(doc, form_data)
            self.generate_circumstances_section(doc, form_data)
            self.generate_summary_section(doc, form_data)
            self.generate_signature(doc)
            
            # Save the document
            return self.save_document(doc)

        except Exception as e:
            print(f"Detailed error in generate: {str(e)}")
            raise

    def save_document(self, doc):
        """Save the generated document."""
        try:
            report_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word Documents", "*.docx")],
                initialfile=f"דוח חקירה_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            )
            if report_path:
                doc.save(report_path)
                return True
            return False

        except Exception as e:
            print(f"Error in save_document: {str(e)}")
            raise