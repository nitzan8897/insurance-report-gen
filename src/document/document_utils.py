from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime

class DocumentUtils:
    def set_document_rtl(self, doc):
        """Sets the whole document to RTL and aligns text right"""
        # Set RTL for section
        section = doc.sections[0]._sectPr
        bidi = OxmlElement('w:bidi')
        section.append(bidi)
        
        # Set RTL for paragraphs
        for paragraph in doc.paragraphs:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            pPr = paragraph._p.get_or_add_pPr()
            bidi = OxmlElement('w:bidi')
            pPr.append(bidi)
            
            # Set RTL for runs
            for run in paragraph.runs:
                rPr = run._r.get_or_add_rPr()
                rtl = OxmlElement('w:rtl')
                rPr.append(rtl)

    def set_run_rtl(self, run):
        """Set Hebrew RTL text properties for a run"""
        run.font.name = 'David'
        run.font.size = Pt(11)
        run.font.rtl = True
        
        # Add RTL paragraph direction
        paragraph = run.element.getparent()
        if paragraph is not None:
            pPr = paragraph.get_or_add_pPr()
            bidi = OxmlElement('w:bidi')
            pPr.append(bidi)
            
            # Set paragraph direction RTL
            if not pPr.xpath('w:textDirection'):
                textDirection = OxmlElement('w:textDirection')
                textDirection.set(qn('w:val'), 'rtl')
                pPr.append(textDirection)

    def add_custom_header_footer(self, doc):
        """
        Add professional header and footer with company branding.
        Simple, clean design suitable for official investigation reports.
        """
        section = doc.sections[0]

        # Add RTL settings to section
        pPr = section._sectPr
        bidi = OxmlElement('w:bidi')
        pPr.append(bidi)

        # Create simple professional header
        header = section.header
        if not header.paragraphs:
            header_para = header.add_paragraph()
        else:
            header_para = header.paragraphs[0]

        header_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        header_run = header_para.add_run('אניגמה חקירות  |  ENIGMA INVESTIGATIONS')
        header_run.font.name = 'David'
        header_run.font.size = Pt(12)
        header_run.bold = True

        # Create simple professional footer
        footer = section.footer
        if not footer.paragraphs:
            footer_para = footer.add_paragraph()
        else:
            footer_para = footer.paragraphs[0]

        footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Contact information in footer
        footer_run = footer_para.add_run(
            'רחוב בית הלל 28, תל-אביב 67017  |  '
            'טל: 03-5222766  |  פקס: 03-5244788  |  '
            'Email: enigmainv@012.net.il'
        )
        footer_run.font.name = 'David'
        footer_run.font.size = Pt(9)

    def make_hebrew_paragraph(self, doc, text, bold=False, size=11, alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY):
        """Create a hebrew paragraph with proper RTL settings - using JUSTIFY like example.docx"""
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(text)
        run.font.name = 'David'
        run.font.size = Pt(size)
        run.font.bold = bold
        paragraph.alignment = alignment
        
        # Set RTL for paragraph
        pPr = paragraph._p.get_or_add_pPr()
        bidi = OxmlElement('w:bidi')
        pPr.append(bidi)
        
        # Set RTL for run
        rPr = run._r.get_or_add_rPr()
        rtl = OxmlElement('w:rtl')
        rPr.append(rtl)
        
        return paragraph

    def add_table_row(self, table, cells_data, bold=False, alignment=WD_PARAGRAPH_ALIGNMENT.RIGHT):
        """Add a row to a table with proper RTL settings"""
        row = table.add_row()
        for i, text in enumerate(cells_data):
            cell = row.cells[i]
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run(text)
            run.font.name = 'David'
            run.font.size = Pt(11)
            run.font.bold = bold
            paragraph.alignment = alignment
            
            # Set RTL for paragraph
            pPr = paragraph._p.get_or_add_pPr()
            bidi = OxmlElement('w:bidi')
            pPr.append(bidi)
            
            # Set RTL for run
            rPr = run._r.get_or_add_rPr()
            rtl = OxmlElement('w:rtl')
            rPr.append(rtl)

    def create_section_header(self, doc, text, level=1):
        """Create a section header with proper formatting"""
        sizes = {1: 14, 2: 12, 3: 11}
        size = sizes.get(level, 11)
        
        paragraph = self.make_hebrew_paragraph(doc, text, bold=True, size=size)
        if level == 1:
            # Add underline for main sections
            run = paragraph.runs[0]
            run.underline = True
        
        return paragraph

    def add_bullet_point(self, doc, text, level=0):
        """Add a bullet point with proper RTL formatting - using JUSTIFY like example.docx"""
        # Don't use style, just create paragraph manually with bullet character
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        # Add bullet character manually (Hebrew-compatible)
        run = paragraph.add_run('• ' + text)
        self.set_run_rtl(run)

        # Add indentation based on level
        paragraph.paragraph_format.left_indent = Inches(0.5 * (level + 1))

        # Set RTL for paragraph
        pPr = paragraph._p.get_or_add_pPr()
        bidi = OxmlElement('w:bidi')
        pPr.append(bidi)

        return paragraph