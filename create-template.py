import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from tkcalendar import DateEntry
from datetime import datetime
import json
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

class InsuranceReportGenerator:
    def __init__(self, root):
        self.root = root
        self.root.title("מערכת דוחות ביטוח")
        self.root.geometry("800x600")
        self.root.configure(bg='#f0f0f0')

        # Data lists
        self.event_types = [
            'גניבת רכב',
            'צד ג\' - רכב',
            'נזק לרכב',
            'חבויות',
            'נח"ל',
            'פריצה לעסק',
            'נזקי אש',
            'פריצה לדירה',
            'נזקי פריצה',
            'גניבת כלי צמ"ה',
            'אבדן תכשיט',
            'גניבת תכשיטים',
            'נזקי מים'
        ]

        self.car_manufacturers = [
            'טויוטה', 'פולקסווגן', 'פורד', 'הונדה', 'שברולט',
            'ניסן', 'יונדאי', 'קיה', 'מרצדס-בנץ', 'ב.מ.וו',
            'אאודי', 'סוזוקי', 'פיג\'ו', 'מאזדה', 'סובארו',
            'סיטרואן', 'מיצובישי', 'רנו', 'סקודה', 'פיאט',
            'לקסוס', 'וולוו', 'ג\'יפ', 'שיווי', 'פורשה',
            'דאצ\'יה', 'כריסלר', 'אופל', 'לנד רובר', 'סיאט',
            'מיני', 'אלפא רומיאו', 'דודג\'', 'קדילק', 'יגואר',
            'סמארט', 'אינפיניטי', 'לינקולן', 'מזראטי', 'טסלה'
        ]

        self.car_colors = [
            'לבן', 'שחור', 'אפור', 'כסף', 'כחול',
            'אדום', 'חום', 'ירוק', 'בז\'', 'זהב',
            'כתום', 'צהוב', 'סגול', 'ורוד', 'תכלת'
        ]

        # Configure Hebrew RTL
        self.root.tk.call('tk', 'scaling', 1.3)

        # Create main notebook for tabs
        self.notebook = ttk.Notebook(root)
        self.notebook.pack(fill='both', expand=True, padx=10, pady=5)

        # Initialize form data
        self.form_data = {}
        self.uploaded_video = None
        self.uploaded_image = None

        # Create tabs
        self.create_tabs()

        # Create bottom buttons
        self.create_bottom_buttons()

        # Load saved data if exists
        self.load_saved_data()

        # Bind event type change
        if 'event_type' in self.form_data:
            self.form_data['event_type'].bind('<<ComboboxSelected>>', self.on_event_type_change)

    def create_tabs(self):
        # Basic Info Tab
        self.basic_frame = ttk.Frame(self.notebook)
        self.notebook.add(self.basic_frame, text='פרטים בסיסיים')
        self.create_basic_info_fields(self.basic_frame)

        # Vehicle Info Tab
        self.vehicle_frame = ttk.Frame(self.notebook)
        self.create_vehicle_info_fields(self.vehicle_frame)

        # Third Party Tab
        self.third_party_frame = ttk.Frame(self.notebook)
        self.create_third_party_fields(self.third_party_frame)

        # Additional Info Tab
        self.additional_frame = ttk.Frame(self.notebook)
        self.notebook.add(self.additional_frame, text='פרטים נוספים')
        self.create_additional_info_fields(self.additional_frame)

    def on_event_type_change(self, event=None):
        event_type = self.form_data['event_type'].get()
        vehicle_related_events = ['גניבת רכב', 'צד ג\' - רכב', 'נזק לרכב', 'נזקי פריצה']

        # Remove all optional tabs
        for tab in [self.vehicle_frame, self.third_party_frame]:
            if tab in self.notebook.tabs():
                self.notebook.forget(tab)

        # Add relevant tabs based on event type
        if event_type in vehicle_related_events:
            self.notebook.add(self.vehicle_frame, text='פרטי רכב')
            if event_type == 'צד ג\' - רכב':
                self.notebook.add(self.third_party_frame, text='פרטי צד ג׳')

        # Update the Additional Info tab based on event type
        self.update_additional_info_fields(event_type)

    def create_basic_info_fields(self, parent):
        canvas = tk.Canvas(parent)
        scrollbar = ttk.Scrollbar(parent, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas)

        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )

        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)

        # Event Type (Dropdown)
        ttk.Label(scrollable_frame, text='סוג תיק').grid(row=0, column=1, padx=5, pady=5, sticky='e')
        event_type = ttk.Combobox(scrollable_frame, values=self.event_types, width=37, state='readonly')
        event_type.grid(row=0, column=0, padx=5, pady=5)
        self.form_data['event_type'] = event_type

        # Date (DatePicker)
        ttk.Label(scrollable_frame, text='תאריך אירוע').grid(row=1, column=1, padx=5, pady=5, sticky='e')
        date_picker = DateEntry(scrollable_frame, width=37, background='darkblue',
                              foreground='white', borderwidth=2, date_pattern='dd/mm/yyyy')
        date_picker.grid(row=1, column=0, padx=5, pady=5)
        self.form_data['event_date'] = date_picker

        # Other basic fields
        other_fields = [
            ('claim_number', 'מספר תביעה'),
            ('full_name', 'שם מלא של המבוטח'),
            ('policy_number', 'מספר פוליסה')
        ]

        for i, (field_name, label_text) in enumerate(other_fields, start=2):
            ttk.Label(scrollable_frame, text=label_text).grid(row=i, column=1, padx=5, pady=5, sticky='e')
            entry = ttk.Entry(scrollable_frame, width=40)
            entry.grid(row=i, column=0, padx=5, pady=5)
            self.form_data[field_name] = entry

        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

    def create_vehicle_info_fields(self, parent):
        canvas = tk.Canvas(parent)
        scrollbar = ttk.Scrollbar(parent, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas)

        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )

        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)

        # Manufacturer (Dropdown)
        ttk.Label(scrollable_frame, text='יצרן הרכב').grid(row=0, column=1, padx=5, pady=5, sticky='e')
        manufacturer = ttk.Combobox(scrollable_frame, values=self.car_manufacturers, width=37, state='readonly')
        manufacturer.grid(row=0, column=0, padx=5, pady=5)
        self.form_data['vehicle_company'] = manufacturer

        # Color (Dropdown)
        ttk.Label(scrollable_frame, text='צבע הרכב').grid(row=1, column=1, padx=5, pady=5, sticky='e')
        color = ttk.Combobox(scrollable_frame, values=self.car_colors, width=37, state='readonly')
        color.grid(row=1, column=0, padx=5, pady=5)
        self.form_data['vehicle_color'] = color

        # Last test date (DatePicker)
        ttk.Label(scrollable_frame, text='תאריך טסט אחרון').grid(row=2, column=1, padx=5, pady=5, sticky='e')
        test_date = DateEntry(scrollable_frame, width=37, background='darkblue',
                            foreground='white', borderwidth=2, date_pattern='dd/mm/yyyy')
        test_date.grid(row=2, column=0, padx=5, pady=5)
        self.form_data['last_roadworthiness_test_date'] = test_date

        # Other vehicle fields
        other_fields = [
            ('vehicle_model', 'דגם הרכב'),
            ('vehicle_manufacture_year', 'שנת ייצור'),
            ('vehicle_license_number', 'מספר רישוי'),
            ('vehicle_engine_type', 'סוג מנוע'),
            ('vehicle_engine_capacity', 'נפח מנוע'),
            ('vehicle_engine_power', 'הספק מנוע'),
            ('vehicle_gearbox', 'סוג גיר')
        ]

        for i, (field_name, label_text) in enumerate(other_fields, start=3):
            ttk.Label(scrollable_frame, text=label_text).grid(row=i, column=1, padx=5, pady=5, sticky='e')
            entry = ttk.Entry(scrollable_frame, width=40)
            entry.grid(row=i, column=0, padx=5, pady=5)
            self.form_data[field_name] = entry

        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

    def create_third_party_fields(self, parent):
        canvas = tk.Canvas(parent)
        scrollbar = ttk.Scrollbar(parent, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas)

        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )

        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)

        # Third-party fields
        other_fields = [
            ('third_party_name', 'שם צד ג\''),
            ('third_party_policy_number', 'מספר פוליסה צד ג\''),
            ('third_party_contact', 'טלפון צד ג\'')
        ]

        for i, (field_name, label_text) in enumerate(other_fields, start=0):
            ttk.Label(scrollable_frame, text=label_text).grid(row=i, column=1, padx=5, pady=5, sticky='e')
            entry = ttk.Entry(scrollable_frame, width=40)
            entry.grid(row=i, column=0, padx=5, pady=5)
            self.form_data[field_name] = entry

        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

    def create_additional_info_fields(self, parent):
        canvas = tk.Canvas(parent)
        scrollbar = ttk.Scrollbar(parent, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas)

        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )

        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)

        # Basic additional fields
        ttk.Label(scrollable_frame, text='נסיבות האירוע').grid(row=0, column=1, padx=5, pady=5, sticky='e')
        circumstances_text = tk.Text(scrollable_frame, width=40, height=5)
        circumstances_text.grid(row=0, column=0, padx=5, pady=5)
        self.form_data['circumstances'] = circumstances_text

        ttk.Label(scrollable_frame, text='סרטוני וידאו').grid(row=1, column=1, padx=5, pady=5, sticky='e')
        video_upload_button = ttk.Button(scrollable_frame, text='בחר קובץ', command=self.upload_video)
        video_upload_button.grid(row=1, column=0, padx=5, pady=5)
        self.form_data['video_file'] = None

        ttk.Label(scrollable_frame, text='התכתבויות').grid(row=2, column=1, padx=5, pady=5, sticky='e')
        image_upload_button = ttk.Button(scrollable_frame, text='בחר קובץ', command=self.upload_image)
        image_upload_button.grid(row=2, column=0, padx=5, pady=5)
        self.form_data['correspondence_image'] = None

        ttk.Label(scrollable_frame, text='החקירה עצמה').grid(row=3, column=1, padx=5, pady=5, sticky='e')
        investigation_text = tk.Text(scrollable_frame, width=40, height=5)
        investigation_text.grid(row=3, column=0, padx=5, pady=5)
        self.form_data['investigation'] = investigation_text

        ttk.Label(scrollable_frame, text='סיכום').grid(row=4, column=1, padx=5, pady=5, sticky='e')
        summary_text = tk.Text(scrollable_frame, width=40, height=5)
        summary_text.grid(row=4, column=0, padx=5, pady=5)
        self.form_data['summary'] = summary_text

        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

    def update_additional_info_fields(self, event_type):
        # Clear existing additional fields
        for widget in self.additional_frame.winfo_children():
            widget.destroy()

        self.create_additional_info_fields(self.additional_frame)  # Recreate basic fields

        # Add vehicle-related fields if applicable
        vehicle_related_events = ['גניבת רכב', 'צד ג\' - רכב', 'נזק לרכב', 'נזקי פריצה']
        if event_type in vehicle_related_events:
            ttk.Label(self.additional_frame, text='פרטי רכב נוספים').grid(row=5, column=0, padx=5, pady=5, sticky='w')
            additional_info_text = tk.Text(self.additional_frame, width=40, height=5)
            additional_info_text.grid(row=6, column=0, padx=5, pady=5)
            self.form_data['vehicle_additional_info'] = additional_info_text

    def create_bottom_buttons(self):
        button_frame = ttk.Frame(self.root)
        button_frame.pack(pady=10)

        save_button = ttk.Button(button_frame, text="שמור", command=self.save_data)
        save_button.pack(side="left", padx=5)

        generate_report_button = ttk.Button(button_frame, text="צור דוח", command=self.generate_report)
        generate_report_button.pack(side="left", padx=5)

    def load_saved_data(self):
        if os.path.exists('saved_data.json'):
            try:
                with open('saved_data.json', 'r', encoding='utf-8') as f:
                    saved_data = json.load(f)
                    for key, value in saved_data.items():
                        if key in self.form_data:
                            if isinstance(self.form_data[key], tk.Text):
                                self.form_data[key].delete('1.0', tk.END)
                                self.form_data[key].insert('1.0', value)
                            elif isinstance(self.form_data[key], (ttk.Entry, ttk.Combobox)):
                                self.form_data[key].delete(0, tk.END)
                                self.form_data[key].insert(0, value)
                            elif isinstance(self.form_data[key], DateEntry):
                                # Handle date entry fields
                                try:
                                    self.form_data[key].set_date(value)
                                except:
                                    pass
                        # Handle file paths separately
                        elif key == 'video_file':
                            self.uploaded_video = value
                        elif key == 'correspondence_image':
                            self.uploaded_image = value
            except Exception as e:
                print(f"Error loading saved data: {str(e)}")

    def upload_video(self):
        file_path = filedialog.askopenfilename(title="בחר קובץ וידאו", filetypes=[("Video Files", "*.mp4;*.avi;*.mov")])
        if file_path:
            self.uploaded_video = file_path
            messagebox.showinfo("Uploaded", "וידאו הועלה בהצלחה!")

    def upload_image(self):
        file_path = filedialog.askopenfilename(title="בחר קובץ תמונה", filetypes=[("Image Files", "*.png;*.jpg;*.jpeg;*.gif")])
        if file_path:
            self.uploaded_image = file_path
            messagebox.showinfo("Uploaded", "תמונה הועלתה בהצלחה!")

    def save_data(self):
        # Save data to a JSON file
        data_to_save = {key: (entry.get() if isinstance(entry, tk.Entry) else entry.get("1.0", tk.END).strip())
                         for key, entry in self.form_data.items() if isinstance(entry, (tk.Entry, tk.Text))}
        
        if self.uploaded_video:
            data_to_save['video_file'] = self.uploaded_video
        if self.uploaded_image:
            data_to_save['correspondence_image'] = self.uploaded_image

        with open('saved_data.json', 'w', encoding='utf-8') as f:
            json.dump(data_to_save, f, ensure_ascii=False, indent=4)

        messagebox.showinfo("Success", "Data saved successfully!")

    def generate_report(self):
        try:
            doc = Document()
            
            # Set document RTL
            self.set_document_rtl(doc)
            # Add custom header and footer
            self.add_custom_header_footer(doc)
            
            # Configure document style
            section = doc.sections[0]
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

            # Add current date and reference number
            date_para = doc.add_paragraph()
            date_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            current_date = datetime.now().strftime("%d ביוני %Y")
            ref_number = datetime.now().strftime("%d%m%y")
            date_run = date_para.add_run(f"{current_date} מספרנו {ref_number}")
            date_run.bold = True
            self.set_run_rtl(date_run)

            # Add "לכבוד" section
            header = doc.add_paragraph()
            header.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            header_run = header.add_run("לכבוד")
            header_run.bold = True
            self.set_run_rtl(header_run)

            company = doc.add_paragraph()
            company.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            company_run = company.add_run("הפניקס -- השכרת רכב בע\"מ")
            company_run.bold = True
            self.set_run_rtl(company_run)

            # Add main header section
            headers = [
                "הנדון : דו\"ח חקירה",
                "======================",
                f"המבוטח : {self.get_field_value('full_name')}",
                f"האירוע : {self.get_field_value('event_type')}",
                f"מספר רישוי מבוטח : {self.get_field_value('vehicle_license_number')}",
                f"מס' רישוי צד ג' : {self.get_field_value('third_party_license_number')}",
                f"תאריך אירוע : {self.get_field_value('event_date')}",
                f"מספר תביעה : {self.get_field_value('claim_number')}",
                "========================="
            ]

            for text in headers:
                para = doc.add_paragraph()
                para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = para.add_run(text)
                run.bold = True
                self.set_run_rtl(run)

            # Add empty paragraph for spacing
            doc.add_paragraph()

            # Add general section (כללי)
            general_title = doc.add_paragraph()
            general_title.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            title_run = general_title.add_run("1. כללי")
            title_run.bold = True
            self.set_run_rtl(title_run)

            # Add general content
            general_text = f"נתבקשנו על ידי חברתכם לבצע חקירה בעקבות הודעת המבוטח על {self.get_field_value('event_type')} שארע/ה לו ברכבו מסוג: {self.get_field_value('vehicle_company')} {self.get_field_value('vehicle_model')} בצבע {self.get_field_value('vehicle_color')}, משנת ייצור {self.get_field_value('vehicle_manufacture_year')}."
            general_para = doc.add_paragraph()
            general_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            general_run = general_para.add_run(general_text)
            self.set_run_rtl(general_run)

            # Add investigation actions
            actions_text = "במסגרת החקירה ביצענו את הפעולות הבאות:"
            actions_para = doc.add_paragraph()
            actions_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            actions_run = actions_para.add_run(actions_text)
            self.set_run_rtl(actions_run)

            # Add bullet points for investigation actions
            bullets = [
                f"פגשנו וחקרנו את נהג הרכב המבוטח -- {self.get_field_value('full_name')} - חקרנו אותו אודות הרכב שבנדון, פרטי התאונה, נסיבותיה ונזקיה.",
                "ערכנו בדיקה במרשתת לאיתור פרסומים רלוונטיים.",
                "ערכנו בדיקה אודות עברו הביטוחי של הרכב."
            ]

            for bullet in bullets:
                bullet_para = doc.add_paragraph(style='List Bullet')
                bullet_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                bullet_run = bullet_para.add_run(bullet)
                self.set_run_rtl(bullet_run)

            # Add findings header
            findings_para = doc.add_paragraph()
            findings_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            findings_run = findings_para.add_run("להלן יובאו ממצאינו:")
            findings_run.bold = True
            self.set_run_rtl(findings_run)

            # Add vehicle details section
            vehicle_title = doc.add_paragraph()
            vehicle_title.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            vehicle_run = vehicle_title.add_run("2. פרטי הרכב")
            vehicle_run.bold = True
            self.set_run_rtl(vehicle_run)

            # Add vehicle details
            vehicle_details = [
                f"יצרן ודגם: {self.get_field_value('vehicle_company')} {self.get_field_value('vehicle_model')}",
                f"צבע: {self.get_field_value('vehicle_color')}",
                f"שנת ייצור: {self.get_field_value('vehicle_manufacture_year')}",
                f"מספר רישוי: {self.get_field_value('vehicle_license_number')}",
                f"נפח מנוע: {self.get_field_value('vehicle_engine_capacity')} סמ\"ק",
                f"סוג תיבת הילוכים: {self.get_field_value('vehicle_gearbox')}"
            ]

            for detail in vehicle_details:
                detail_para = doc.add_paragraph()
                detail_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                detail_para.paragraph_format.left_indent = Pt(20)
                detail_run = detail_para.add_run(detail)
                self.set_run_rtl(detail_run)

            # Add circumstances section
            if self.get_field_value('circumstances'):
                circum_title = doc.add_paragraph()
                circum_title.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                circum_run = circum_title.add_run("3. נסיבות האירוע")
                circum_run.bold = True
                self.set_run_rtl(circum_run)

                circum_para = doc.add_paragraph()
                circum_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                circum_run = circum_para.add_run(self.get_field_value('circumstances'))
                self.set_run_rtl(circum_run)

            # Add summary section
            if self.get_field_value('summary'):
                summary_title = doc.add_paragraph()
                summary_title.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                summary_run = summary_title.add_run("4. סיכום")
                summary_run.bold = True
                self.set_run_rtl(summary_run)

                summary_para = doc.add_paragraph()
                summary_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                summary_run = summary_para.add_run(self.get_field_value('summary'))
                self.set_run_rtl(summary_run)

            # Add signature
            doc.add_paragraph()  # Add spacing
            signature_para = doc.add_paragraph()
            signature_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            signature_run = signature_para.add_run("בכבוד רב,")
            self.set_run_rtl(signature_run)

            company_para = doc.add_paragraph()
            company_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            company_run = company_para.add_run("אניגמה חקירות")
            self.set_run_rtl(company_run)

            # Save the document
            report_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word Documents", "*.docx")],
                initialfile=f"דוח חקירה_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            )
            if report_path:
                doc.save(report_path)
                messagebox.showinfo("הצלחה", "הדו\"ח נוצר בהצלחה!")

        except Exception as e:
            messagebox.showerror("שגיאה", f"אירעה שגיאה בעת יצירת הדו\"ח: {str(e)}")

    def set_run_rtl(self, run):
        """Set Hebrew RTL text properties"""
        run.font.name = 'David'
        run.font.size = Pt(11)
        run.font.rtl = True
        
        # Add RTL paragraph direction
        paragraph = run.element.getparent()
        if paragraph is not None:
            pPr = paragraph.get_or_add_pPr()
            bidi = OxmlElement('w:bidi')
            pPr.append(bidi)
            
            # Set paragraph direction RTL
            if not pPr.xpath('w:textDirection'):
                textDirection = OxmlElement('w:textDirection')
                textDirection.set(qn('w:val'), 'rtl')
                pPr.append(textDirection)

    def get_field_value(self, field_name):
        """Get value from form field"""
        if field_name in self.form_data:
            field = self.form_data[field_name]
            if isinstance(field, tk.Text):
                return field.get("1.0", tk.END).strip()
            elif isinstance(field, DateEntry):
                return field.get_date().strftime("%d.%m.%Y")
            elif isinstance(field, ttk.Combobox):
                return field.get()
            else:
                return field.get()
        return ""
    def set_rtl(self, paragraph):
        """Set paragraph direction to RTL"""
        pPr = paragraph._p.get_or_add_pPr()
        bidi = OxmlElement('w:bidi')
        pPr.append(bidi)

    def set_document_rtl(self, doc):
        """Set the whole document to RTL and align text right"""
        # Set RTL for section
        section = doc.sections[0]._sectPr
        bidi = OxmlElement('w:bidi')
        section.append(bidi)
        
        # Set RTL for paragraphs
        for paragraph in doc.paragraphs:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            pPr = paragraph._p.get_or_add_pPr()
            bidi = OxmlElement('w:bidi')
            pPr.append(bidi)
            
            # Set RTL for runs
            for run in paragraph.runs:
                rPr = run._r.get_or_add_rPr()
                rtl = OxmlElement('w:rtl')
                rPr.append(rtl)

    def make_hebrew_paragraph(self, doc, text, bold=False, size=11, alignment=WD_PARAGRAPH_ALIGNMENT.RIGHT):
        """Create a hebrew paragraph with proper RTL settings"""
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(text)
        run.font.name = 'David'
        run.font.size = Pt(size)
        run.font.bold = bold
        paragraph.alignment = alignment
        
        # Set RTL for paragraph
        pPr = paragraph._p.get_or_add_pPr()
        bidi = OxmlElement('w:bidi')
        pPr.append(bidi)
        
        # Set RTL for run
        rPr = run._r.get_or_add_rPr()
        rtl = OxmlElement('w:rtl')
        rPr.append(rtl)
        
        return paragraph

    def add_headers_footers(self, doc):
        """Add headers and footers to the document"""
        # Header
        header = doc.sections[0].header
        header_para = header.paragraphs[0]
        header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        current_date = datetime.now().strftime("%d ביוני %Y")
        ref_number = datetime.now().strftime("%d%m%y")
        header_run = header_para.add_run(f"{current_date} מספרנו {ref_number}")
        header_run.font.name = 'David'
        header_run.font.size = Pt(11)
        header_run.bold = True
        self.set_rtl(header_para)

        # Footer
        footer = doc.sections[0].footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_text = """הודעת המבוטח
        הודעת בעל רכב צ"ג
        הקלטות חקירות ושיחות טלפון"""
        footer_run = footer_para.add_run(footer_text)
        footer_run.font.name = 'David'
        footer_run.font.size = Pt(11)
        footer_run.bold = True
        self.set_rtl(footer_para)
        
    def add_custom_header_footer(self, doc):
        # Header
        section = doc.sections[0]
        
        # Add RTL settings to section
        pPr = section._sectPr
        bidi = OxmlElement('w:bidi')
        pPr.append(bidi)
        
        # Create header
        header = section.header
        
        # Create 2-column table for header
        width = section.page_width - (section.left_margin + section.right_margin)
        header_table = header.add_table(1, 2, width)
        
        # English header cell
        eng_cell = header_table.cell(0, 0)
        eng_para = eng_cell.paragraphs[0]
        eng_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        eng_text = "ENIGMA INVESTIGATIONS"
        for i, char in enumerate(eng_text):
            run = eng_para.add_run(char)
            run.font.name = "Aharoni"
            run.font.size = Pt(14)
            color_value = int(255 * (1 - (i / len(eng_text))))
            run.font.color.rgb = RGBColor(color_value, 0, 0)

        # Hebrew header cell
        heb_cell = header_table.cell(0, 1)
        heb_para = heb_cell.paragraphs[0]
        heb_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        heb_text = "אניגמה חקירות"
        for i, char in enumerate(heb_text):
            run = heb_para.add_run(char)
            run.font.name = "Aharoni"
            run.font.size = Pt(14)
            color_value = int(255 * (1 - (i / len(heb_text))))
            run.font.color.rgb = RGBColor(color_value, 0, 0)

        # Footer
        footer = section.footer
        
        # Hebrew address
        heb_address = footer.add_paragraph()
        heb_address.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        heb_run = heb_address.add_run('רחוב בית הלל 28, ת"א 67017   טל: 03-5222766  פקס: 03-5244788')
        heb_run.font.name = 'David'
        heb_run.font.size = Pt(8)

        # English address
        eng_address = footer.add_paragraph()
        eng_address.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        eng_run = eng_address.add_run('E-mail: enigmainv@012.net.il                               28 Beit Hillel St. Tel-Aviv 67017  Tel: 03-5222766  Fax: 03-5244788')
        eng_run.font.name = 'Arial'
        eng_run.font.size = Pt(8)

        # First separator
        sep1 = footer.add_paragraph()
        sep1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        sep1.add_run('** **').bold = True

        # Disclaimer
        disclaimer = footer.add_paragraph()
        disclaimer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        disc_run = disclaimer.add_run('כל האינפורמציה הנמסרת על ידינו מיועדת באופן סודי ביותר, אך ורק עבור המזמין לבדו והמזמין יהיה אחראי לכל תוצאה ו/או נזק העלולים להיגרם לנו מחמת האינפורמציה או גילויה לאחרים. האינפורמציה היא רכושנו הבלעדי ויש להחזירה לפי דרישתנו.')
        disc_run.font.name = 'David'
        disc_run.font.size = Pt(8)

        # Second separator
        sep2 = footer.add_paragraph()
        sep2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        sep2.add_run('** **').bold = True

def main():
    root = tk.Tk()
    app = InsuranceReportGenerator(root)
    root.mainloop()

if __name__ == "__main__":
    main()
